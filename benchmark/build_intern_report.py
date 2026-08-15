from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
REPO_ROOT = ROOT.parent
RUN_DIR = ROOT / "results" / "acceptance-test-20260724"
OUTPUT_DIR = REPO_ROOT / "outputs" / "019f91e9-23df-7233-b7ee-e04f24dbc4b5"
OUTPUT_PATH = OUTPUT_DIR / "标书审核基准测试报告_实习生版.docx"
CHART_PATH = OUTPUT_DIR / "intern_report_metrics.png"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
LIGHT_RED = "FCE4D6"
LIGHT_GREEN = "E2F0D9"
GOLD = "FFF2CC"
RED = "9C0006"
GREEN = "548235"
DARK = "202124"
MUTED = "666666"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


metrics = load_json(RUN_DIR / "metrics.json")
summary = load_json(RUN_DIR / "summary.json")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"表格列宽总和必须为9360，实际为{sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=11, bold=None, color=DARK, italic=None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, after=6, before=0, line_spacing=1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_body(doc, text: str, *, bold_lead: str | None = None, after=6) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.clear if False else None


def add_labeled_item(doc, label: str, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, after=5, line_spacing=1.12)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_run_font(r2)


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE, color=NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line_spacing=1.12)
    r1 = p.add_run(label + " ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def format_table_text(table, header=True, font_size=9.5) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                style_paragraph(p, after=0, line_spacing=1.08)
                for run in p.runs:
                    set_run_font(
                        run,
                        size=font_size,
                        bold=(row_index == 0 and header),
                        color=("FFFFFF" if row_index == 0 and header else DARK),
                    )
        if row_index == 0 and header:
            set_repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell, NAVY)


def add_metric_chart() -> None:
    image = Image.new("RGB", (1760, 560), "white")
    draw = ImageDraw.Draw(image)
    font_path = r"C:\Windows\Fonts\simhei.ttf"
    title_font = ImageFont.truetype(font_path, 36)
    label_font = ImageFont.truetype(font_path, 27)
    small_font = ImageFont.truetype(font_path, 23)
    value_font = ImageFont.truetype(font_path, 25)

    draw.rounded_rectangle((20, 20, 855, 540), radius=18, outline="#D9E2F3", width=3)
    draw.rounded_rectangle((905, 20, 1740, 540), radius=18, outline="#D9E2F3", width=3)
    draw.text((270, 45), "上线门槛对比", font=title_font, fill="#17365D")
    draw.text((1130, 45), "30个问题最终去了哪里", font=title_font, fill="#17365D")

    actual = [50, 90]
    target = [80, 95]
    names = ["F1", "重大问题召回率"]
    baseline = 470
    chart_height = 320
    for index, name in enumerate(names):
        center = 245 + index * 340
        actual_height = int(chart_height * actual[index] / 100)
        target_height = int(chart_height * target[index] / 100)
        draw.rectangle(
            (center - 75, baseline - actual_height, center - 10, baseline),
            fill="#C0504D",
        )
        draw.rectangle(
            (center + 20, baseline - target_height, center + 85, baseline),
            fill="#4F81BD",
        )
        draw.text((center - 72, baseline - actual_height - 34), f"{actual[index]}%", font=value_font, fill="#9C0006")
        draw.text((center + 23, baseline - target_height - 34), f"{target[index]}%", font=value_font, fill="#17365D")
        name_width = draw.textbbox((0, 0), name, font=small_font)[2]
        draw.text((center - name_width // 2, 485), name, font=small_font, fill="#333333")
    draw.rectangle((590, 105, 620, 135), fill="#C0504D")
    draw.text((630, 105), "实测", font=small_font, fill="#333333")
    draw.rectangle((700, 105, 730, 135), fill="#4F81BD")
    draw.text((740, 105), "门槛", font=small_font, fill="#333333")

    labels = [("标准答案", 30, "#4F81BD"), ("AI命中", 10, "#70AD47"), ("AI漏检", 20, "#C0504D")]
    max_width = 520
    for index, (label, value, color) in enumerate(labels):
        y = 145 + index * 115
        draw.text((960, y + 10), label, font=label_font, fill="#333333")
        width = int(max_width * value / 30)
        draw.rounded_rectangle((1120, y, 1120 + width, y + 55), radius=8, fill=color)
        draw.text((1135 + width, y + 10), str(value), font=value_font, fill="#333333")
    draw.text((1050, 490), "系统每份只输出1条，平均漏掉2条", font=small_font, fill="#9C0006")
    image.save(CHART_PATH)


OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
add_metric_chart()

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.4)
section.footer_distance = Inches(0.4)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for level, size, color, before, after in (
    (1, 16, BLUE, 16, 8),
    (2, 13, BLUE, 12, 6),
    (3, 11.5, NAVY, 8, 4),
):
    style = styles[f"Heading {level}"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("智能标书审核系统｜基准测试说明")
set_run_font(hr, size=9, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("实习生培训版  ·  第 ")
set_run_font(fr, size=9, color=MUTED)
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
fp._p.append(field)
fr2 = fp.add_run(" 页")
set_run_font(fr2, size=9, color=MUTED)

# Opening masthead
p = doc.add_paragraph()
style_paragraph(p, after=4, before=10)
r = p.add_run("基准测试报告")
set_run_font(r, size=24, bold=True, color=NAVY)

p = doc.add_paragraph()
style_paragraph(p, after=14)
r = p.add_run("智能标书审核系统 · 实习生易读版")
set_run_font(r, size=14, color=MUTED)

meta = [
    ("测试版本", "Silver Benchmark v1.0"),
    ("测试范围", "test 盲测集 10 份标书，注入页模式"),
    ("运行日期", "2026年7月24日"),
    ("上线门槛", "F1 ≥ 80%，Critical Recall ≥ 95%"),
    ("最终结论", "未通过上线门槛"),
]
for label, value in meta:
    p = doc.add_paragraph()
    style_paragraph(p, after=2)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run_font(r2, color=(RED if label == "最终结论" else DARK))

add_callout(
    doc,
    "一句话结论：",
    "系统能发现最明显的重大限制条件，但经常把同一页的三个问题合并成一个，只找出30个标准问题中的10个，因此当前不能作为“重大问题不会漏”的自动审核系统上线。",
    fill=LIGHT_RED,
    color=RED,
)

doc.add_heading("1. 先理解：这次测试到底在做什么", level=1)
add_body(
    doc,
    "可以把基准测试理解成一次闭卷考试。我们提前准备好标准答案，再让AI独立审核，最后由评分程序逐条对答案。"
)
add_labeled_item(doc, "标准答案：", "annotations.jsonl 中的人工设计问题，也叫“真值”或 Gold。")
add_labeled_item(doc, "AI答卷：", "predictions.jsonl 中系统实际输出的风险 finding。")
add_labeled_item(doc, "阅卷程序：", "evaluate.py，负责匹配问题类型、引用原文和页码，并计算各项指标。")
add_labeled_item(doc, "成绩单：", "metrics.json 和 summary.md，告诉我们是否达到上线门槛。")

add_callout(
    doc,
    "重要边界：",
    "测试问题是人工追加到公开标书最后一页的合成问题，不代表原标书真实违法。本轮使用 injected 模式，只验证“已知问题能不能被识别”，还不能完全代表真实生产环境的误报率。",
    fill=GOLD,
    color="7A5A00",
)

doc.add_heading("2. 测试数据由什么组成", level=1)
add_body(
    doc,
    "完整数据集有50份公开招标文件，每份追加3条测试问题，共150条。为了避免一边调系统一边偷看答案，数据被分为训练、验证和盲测三部分。"
)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["数据分组", "文档数", "问题数", "用途"]
for i, text in enumerate(headers):
    table.cell(0, i).text = text
for row in [
    ("Train", "30", "90", "开发和调整提示词、规则"),
    ("Dev", "10", "30", "开发期间反复验证"),
    ("Test", "10", "30", "最终盲测，不应用于调参"),
]:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
set_table_geometry(table, [1500, 1200, 1200, 5460])
format_table_text(table)

doc.add_heading("3. 本次自动验收是怎么跑的", level=1)
for label, text in [
    ("步骤1｜启动引擎：", "一键脚本自动启动Rust审核服务并检查健康状态。"),
    ("步骤2｜上传文件：", "依次上传10份test标书，解析全文并建立条款块。"),
    ("步骤3｜定位测试页：", "通过页码和注入原文定位最后一页，兼容跨页合并的chunk。"),
    ("步骤4｜运行审核：", "调用事实核查、程序合规、规则、语义风险、评分、需求、合同等7类Agent。"),
    ("步骤5｜等待结果：", "每份完成后立即保存，网络中断可断点续跑。"),
    ("步骤6｜自动评分：", "将AI finding与30条真值逐条匹配，输出PASS或FAIL。"),
]:
    add_labeled_item(doc, label, text)

doc.add_page_break()
doc.add_heading("4. 本次基准测试结果", level=1)
add_body(
    doc,
    "10份盲测文档全部执行成功，运行故障为0。系统一共输出10条finding，而标准答案有30条。最终成绩如下。"
)
result_rows = [
    ("Precision", "100%", "无硬性门槛", "输出的10条都能对应到真实注入问题"),
    ("Recall", "33.33%", "越高越好", "30个问题只找到10个，漏掉20个"),
    ("F1", "50%", "≥ 80%", "未通过"),
    ("Critical Recall", "90%", "≥ 95%", "10个重大问题找到9个，漏1个"),
    ("Critical严重度正确率", "0%", "越高越好", "找到的重大问题全部标成High"),
    ("严重度一致率", "10%", "越高越好", "10个命中问题中仅1个等级正确"),
]
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
for i, text in enumerate(["指标", "实测", "门槛/期望", "实习生应该怎样理解"]):
    table.cell(0, i).text = text
for row in result_rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
    if row[0] in {"F1", "Critical Recall", "Critical严重度正确率"}:
        set_cell_shading(cells[1], LIGHT_RED)
set_table_geometry(table, [1900, 1200, 1500, 4760])
format_table_text(table, font_size=9.2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
p.add_run().add_picture(str(CHART_PATH), width=Inches(6.35))

add_callout(
    doc,
    "最终门禁：FAIL",
    "F1只有50%，低于80%；Critical Recall只有90%，低于95%。两个条件必须同时满足，因此当前版本不能通过上线验收。",
    fill=LIGHT_RED,
    color=RED,
)

doc.add_heading("5. 每个指标到底是什么意思", level=1)
doc.add_heading("5.1 Precision：AI报出来的问题有多少是真的", level=2)
add_body(
    doc,
    "本次AI报了10条，其中10条都能对应标准答案，所以Precision是100%。这说明它报出来的内容比较可信，但不能说明它没有漏报。"
)
add_callout(
    doc,
    "容易误解的地方：",
    "Precision 100%不等于系统很好。一个学生只答1道题并且答对，也可以有100%正确率，但其余题目全都没答。",
    fill=GOLD,
    color="7A5A00",
)

doc.add_heading("5.2 Recall：标准答案里的问题找回了多少", level=2)
add_body(
    doc,
    "标准答案有30条，AI只匹配到10条，因此Recall = 10 ÷ 30 = 33.33%。换句话说，每3个问题中平均只发现1个。对审核系统而言，Recall过低意味着漏检很多。"
)

doc.add_heading("5.3 F1：同时看“报得准”和“找得全”", level=2)
add_body(
    doc,
    "F1会综合Precision和Recall。即使Precision很高，只要Recall很低，F1仍会被拉低。本次F1为50%，低于80%的上线要求。"
)

doc.add_heading("5.4 Critical Recall：重大问题有没有漏", level=2)
add_body(
    doc,
    "盲测集中有10个Critical问题，系统命中9个，Critical Recall为90%。门槛是95%，在只有10个重大样本时，漏1个就会从100%直接降到90%，所以本次没有通过。"
)

doc.add_heading("5.5 严重度一致率：发现之后有没有分对等级", level=2)
add_body(
    doc,
    "发现问题只是第一步，还要判断严重程度。系统命中的9个Critical问题全部被标成High，Critical严重度判定正确率为0%。如果前端只重点提醒Critical，这些重大问题可能无法进入最高优先级。"
)

doc.add_page_break()
doc.add_heading("6. 哪些问题识别得好，哪些问题容易漏", level=1)
category_rows = [
    ("Critical", "地域注册限制", "2/2", "较好"),
    ("Critical", "指定品牌", "2/2", "较好"),
    ("Critical", "无关资格条件", "2/2", "较好"),
    ("Critical", "特定区域业绩", "2/2", "较好"),
    ("Critical", "经营规模门槛", "1/2", "漏1条"),
    ("High", "本地奖项加分", "1/2", "漏1条"),
    ("High", "期限不足、保证金、厂家授权、主观评分", "0/8", "全部漏检"),
    ("Medium", "验收、知识产权、单方变更、日期、违约责任", "0/10", "全部漏检"),
]
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
for i, text in enumerate(["等级", "问题类别", "命中/总数", "表现"]):
    table.cell(0, i).text = text
for row in category_rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
    if "0/" in row[2] or "漏" in row[3]:
        set_cell_shading(cells[3], LIGHT_RED)
    else:
        set_cell_shading(cells[3], LIGHT_GREEN)
set_table_geometry(table, [1300, 4760, 1500, 1800])
format_table_text(table, font_size=9.2)

add_body(
    doc,
    "从这个分布可以看出，系统倾向于抓取每页最明显的第一个重大限制条件，却没有把同一条款块里的后续问题拆成独立finding。"
)

doc.add_heading("7. 数据文件分别是干什么的", level=1)
file_rows = [
    ("source_manifest.csv / jsonl", "试卷清单", "记录50份标书来源、页数、文件哈希和数据分组，证明测试材料可追溯。"),
    ("annotations.csv / jsonl", "标准答案", "每一条应被发现的问题，包括类型、等级、页码、原文、原因、建议和法规。"),
    ("taxonomy.csv", "题型字典", "定义15类风险、别名和默认严重度，帮助评分器识别“地域限制”和“地域歧视”是同类表达。"),
    ("mutated/*.pdf", "考试试卷", "在公开标书末页追加3条测试问题后的PDF，实际交给审核系统。"),
    ("predictions.jsonl", "AI答卷", "系统真正输出的finding；本次10份文档共10条。"),
    ("metrics.json", "详细成绩单", "包含TP、FP、FN、各类别Recall、漏检清单和门禁结果。"),
    ("summary.md / summary.json", "验收结论", "给人快速阅读或给程序读取的最终PASS/FAIL摘要。"),
    ("documents/BID-xxx.json", "单份答题记录", "保存每份PDF的解析信息、条款块、耗时和原始finding，便于定位问题。"),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for i, text in enumerate(["文件", "通俗比喻", "它反映什么"]):
    table.cell(0, i).text = text
for row in file_rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
set_table_geometry(table, [2450, 1500, 5410])
format_table_text(table, font_size=8.8)

doc.add_heading("8. 这些结果反映了系统的什么问题", level=1)
add_labeled_item(
    doc,
    "问题一｜finding粒度过粗：",
    "每份测试页有3个问题，系统每份只输出1条。虽然reason中偶尔顺带提到其他风险，但没有形成独立问题、独立等级和独立整改建议，评分时仍属于漏检。",
)
add_labeled_item(
    doc,
    "问题二｜严重度整体偏低：",
    "9个已发现的Critical全部被标为High。系统可能在合并、法条验证或最终分级阶段统一降级。",
)
add_labeled_item(
    doc,
    "问题三｜Agent结果可能被覆盖或合并：",
    "程序、评分和合同类问题几乎全漏，说明相关Agent没有产出，或者产出在Merge、Legal Verify、Triage阶段被合并掉。",
)
add_labeled_item(
    doc,
    "问题四｜当前测试仍是合成注入测试：",
    "它证明系统对明确条款的识别能力不足，但还没有测量真实标书全文的穷尽Precision，因此不能用这次100% Precision宣传真实业务准确率。",
)

doc.add_heading("9. 当前能不能上线", level=1)
add_callout(
    doc,
    "结论：不建议正式上线。",
    "当前版本可以作为内部辅助工具试用，但必须保留人工复核；不能对外承诺“重大问题不会漏”，也不应自动替代政府采购专业人员。",
    fill=LIGHT_RED,
    color=RED,
)
add_body(
    doc,
    "如果只是内部试点，建议在页面上明确标注“AI结果仅供参考”，并要求审核人员逐页确认。等修复后重新跑同一test集，达到F1 ≥ 80%、Critical Recall ≥ 95%，再进入真实原文双人盲标阶段。"
)

doc.add_heading("10. 实习生接下来应该做什么", level=1)
for label, text in [
    ("任务1｜检查输出结构：", "确保同一个chunk里的每个独立风险都生成独立finding，不能把三条风险塞进一条reason。"),
    ("任务2｜检查严重度映射：", "重点跟踪Critical在Agent输出、Merge、Legal Verify、Triage后的等级变化。"),
    ("任务3｜查看漏检清单：", "打开metrics.json的missed数组，按category_code统计和复现。"),
    ("任务4｜每次改动后复跑：", "双击benchmark/run_acceptance.cmd，等待summary.md生成，不要用test答案反向修改提示词。"),
    ("任务5｜记录版本：", "保存模型名称、提示词版本、代码提交和运行编号，保证成绩可以重复。"),
]:
    add_labeled_item(doc, label, text)

add_callout(
    doc,
    "记住三个判断：",
    "Precision看“报得准不准”，Recall看“找得全不全”，Critical Recall看“重大问题漏没漏”。审核系统最怕的不是多提醒一次，而是把真正的重大问题漏掉。",
    fill=LIGHT_BLUE,
    color=NAVY,
)

doc.add_heading("附录：本次验收文件位置", level=1)
for label, text in [
    ("最终报告：", "benchmark/results/acceptance-test-20260724/summary.md"),
    ("完整指标：", "benchmark/results/acceptance-test-20260724/metrics.json"),
    ("AI预测：", "benchmark/results/acceptance-test-20260724/predictions.jsonl"),
    ("一键复跑：", "benchmark/run_acceptance.cmd"),
    ("标准答案：", "benchmark/data/annotations.jsonl"),
]:
    add_labeled_item(doc, label, text)

# Metadata and save
doc.core_properties.title = "标书审核基准测试报告（实习生版）"
doc.core_properties.subject = "Silver Benchmark v1.0 test盲测结果说明"
doc.core_properties.author = "智能标书审核项目组"
doc.core_properties.keywords = "标书审核, 基准测试, F1, Recall, Critical Recall"
doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
